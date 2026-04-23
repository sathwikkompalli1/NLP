"""
IQAS Document Loader
=====================
Ingest PDF, TXT, and DOCX documents with metadata preservation.
"""

from __future__ import annotations

import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Union

from utils.logger import get_logger

log = get_logger("document_loader")


# ──────────────────────────── Data Model ────────────────────────────


@dataclass
class Document:
    """A loaded document page/section with metadata."""
    id: str
    text: str
    source: str
    filename: str
    page_num: Optional[int] = None
    total_pages: Optional[int] = None

    def __post_init__(self):
        if not self.id:
            self.id = str(uuid.uuid4())[:12]


# ──────────────────────────── Document Loader ────────────────────────────


class DocumentLoader:
    """
    Load documents from various file formats.

    Supported formats:
        - PDF (via PyMuPDF/fitz)
        - TXT (plain text)
        - DOCX (via python-docx)
    """

    def load_pdf(self, path: Union[str, Path]) -> List[Document]:
        """
        Load a PDF file, extracting text page by page.

        Args:
            path: Path to the PDF file.

        Returns:
            List of Document objects, one per page.
        """
        import fitz  # PyMuPDF

        path = Path(path)
        if not path.exists():
            log.error(f"PDF file not found: {path}")
            return []

        documents = []
        try:
            pdf = fitz.open(str(path))
            total_pages = len(pdf)
            doc_id = str(uuid.uuid4())[:8]

            for page_num in range(total_pages):
                page = pdf[page_num]
                text = page.get_text("text")

                if text.strip():
                    documents.append(Document(
                        id=f"{doc_id}_p{page_num + 1}",
                        text=text,
                        source=str(path),
                        filename=path.name,
                        page_num=page_num + 1,
                        total_pages=total_pages,
                    ))

            pdf.close()
            log.info(f"Loaded PDF: {path.name} ({total_pages} pages, {len(documents)} with text)")
        except Exception as e:
            log.error(f"Failed to load PDF {path}: {e}")

        return documents

    def load_txt(self, path: Union[str, Path]) -> List[Document]:
        """
        Load a plain text file.

        Args:
            path: Path to the TXT file.

        Returns:
            List containing one Document object.
        """
        path = Path(path)
        if not path.exists():
            log.error(f"TXT file not found: {path}")
            return []

        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            if text.strip():
                doc = Document(
                    id=str(uuid.uuid4())[:8],
                    text=text,
                    source=str(path),
                    filename=path.name,
                    page_num=1,
                    total_pages=1,
                )
                log.info(f"Loaded TXT: {path.name} ({len(text)} chars)")
                return [doc]
            else:
                log.warning(f"TXT file is empty: {path.name}")
                return []
        except Exception as e:
            log.error(f"Failed to load TXT {path}: {e}")
            return []

    def load_docx(self, path: Union[str, Path]) -> List[Document]:
        """
        Load a DOCX file using python-docx.

        Args:
            path: Path to the DOCX file.

        Returns:
            List containing one Document object.
        """
        from docx import Document as DocxDocument

        path = Path(path)
        if not path.exists():
            log.error(f"DOCX file not found: {path}")
            return []

        try:
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)

            if text.strip():
                result = Document(
                    id=str(uuid.uuid4())[:8],
                    text=text,
                    source=str(path),
                    filename=path.name,
                    page_num=1,
                    total_pages=1,
                )
                log.info(f"Loaded DOCX: {path.name} ({len(paragraphs)} paragraphs)")
                return [result]
            else:
                log.warning(f"DOCX file is empty: {path.name}")
                return []
        except Exception as e:
            log.error(f"Failed to load DOCX {path}: {e}")
            return []

    def load_any(self, path: Union[str, Path]) -> List[Document]:
        """
        Auto-detect file format and load accordingly.

        Args:
            path: Path to the document file.

        Returns:
            List of Document objects.
        """
        path = Path(path)
        ext = path.suffix.lower()

        loaders = {
            ".pdf": self.load_pdf,
            ".txt": self.load_txt,
            ".text": self.load_txt,
            ".docx": self.load_docx,
        }

        loader = loaders.get(ext)
        if loader is None:
            log.warning(f"Unsupported file format: {ext} — treating as plain text")
            return self.load_txt(path)

        return loader(path)

    def batch_load(self, paths: List[Union[str, Path]]) -> List[Document]:
        """
        Load multiple documents.

        Args:
            paths: List of file paths.

        Returns:
            Combined list of Document objects from all files.
        """
        all_docs: List[Document] = []
        for path in paths:
            docs = self.load_any(path)
            all_docs.extend(docs)

        log.info(f"Batch loaded {len(all_docs)} document sections from {len(paths)} files")
        return all_docs
